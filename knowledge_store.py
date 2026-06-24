"""
knowledge_store.py - 知识库核心模块

支持：文档加载 → 分块 → Embedding → 向量存储 → RAG 检索

依赖：
    pip install chromadb pymupdf python-docx beautifulsoup4 lxml

Embedding 模型（通过 Ollama）：
    ollama pull nomic-embed-text
"""

import os
import uuid
import logging
import hashlib
import json
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
from dataclasses import dataclass, field
from datetime import datetime

logger = logging.getLogger(__name__)

# ========== 配置 ==========
KB_ROOT = Path("D:/qwen3-asr/knowledge_base")
KB_ROOT.mkdir(exist_ok=True)
CHROMA_PATH = KB_ROOT / "chroma_db"
DOCS_PATH = KB_ROOT / "documents"
INDEX_FILE = KB_ROOT / "index.json"


# ========== 数据结构 ==========

@dataclass
class Document:
    """文档数据结构"""
    doc_id: str
    filename: str
    content: str
    metadata: Dict = field(default_factory=dict)
    chunk_count: int = 0


@dataclass
class Chunk:
    """文本块数据结构"""
    chunk_id: str
    doc_id: str
    text: str
    chunk_index: int
    metadata: Dict = field(default_factory=dict)


@dataclass
class SearchHit:
    """检索结果"""
    chunk_id: str
    text: str
    metadata: Dict
    distance: float
    score: float


# ========== 1. 文档加载器 ==========

class DocumentLoader:
    """支持多格式的文档加载器"""

    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.pdf', '.docx', '.html', '.pptx', '.xlsx'}

    def load(self, file_path: str) -> Document:
        """加载文档"""
        ext = Path(file_path).suffix.lower()

        if ext == '.txt' or ext == '.md':
            return self._load_text(file_path)
        elif ext == '.pdf':
            return self._load_pdf(file_path)
        elif ext == '.docx':
            return self._load_docx(file_path)
        elif ext == '.html':
            return self._load_html(file_path)
        else:
            raise ValueError(f"不支持的文档格式: {ext}")

    def _load_text(self, path: str) -> Document:
        """加载纯文本/Markdown"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self._make_doc(path, content)

    def _load_pdf(self, path: str) -> Document:
        """加载 PDF（使用 PyMuPDF）"""
        import fitz  # pymupdf
        text_parts = []
        with fitz.open(path) as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[第{page.number + 1}页]\n{page_text}")
        content = "\n\n".join(text_parts)
        return self._make_doc(path, content)

    def _load_docx(self, path: str) -> Document:
        """加载 Word 文档"""
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        content = "\n\n".join(paragraphs)
        return self._make_doc(path, content)

    def _load_html(self, path: str) -> Document:
        """加载 HTML（提取正文）"""
        from bs4 import BeautifulSoup
        with open(path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        # 移除脚本和样式
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            tag.decompose()
        # 获取文本
        content = soup.get_text(separator="\n\n", strip=True)
        # 清理多余空行
        content = re.sub(r'\n{3,}', '\n\n', content)
        return self._make_doc(path, content)

    def _make_doc(self, path: str, content: str) -> Document:
        """构建 Document 对象"""
        p = Path(path)
        # 生成唯一 ID
        doc_id = hashlib.md5(f"{p}{p.stat().st_size}".encode()).hexdigest()[:12]
        return Document(
            doc_id=doc_id,
            filename=p.name,
            content=content,
            metadata={
                "source": str(p.absolute()),
                "filename": p.name,
                "extension": p.suffix,
                "size_bytes": p.stat().st_size,
                "loaded_at": datetime.now().isoformat()
            }
        )


# ========== 2. 文本分块器 ==========

class TextChunker:
    """
    智能文本分块器

    策略：
    1. 优先按段落分块（保持语义完整）
    2. 单段落过长时按句子分块
    3. 相邻块之间有重叠（overlap），保证检索连续性
    """

    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        """
        Args:
            chunk_size: 每块最大字符数
            overlap: 相邻块之间的重叠字符数
        """
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, doc: Document) -> List[Chunk]:
        """将文档分块"""
        chunks = []
        content = doc.content

        # 1. 先按段落分割
        paragraphs = re.split(r'\n\s*\n', content)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        current_block = ""
        chunk_index = 0

        for para in paragraphs:
            # 如果单个段落就超过 chunk_size，按句子分割
            if len(para) > self.chunk_size:
                # 先保存当前块
                if current_block:
                    chunks.append(self._make_chunk(doc, current_block, chunk_index))
                    chunk_index += 1
                    current_block = ""

                # 处理超长段落：按句子分割
                sentences = self._split_sentences(para)
                for sent in sentences:
                    if len(current_block) + len(sent) + 1 <= self.chunk_size:
                        current_block = (current_block + "\n" + sent).strip()
                    else:
                        if current_block:
                            chunks.append(self._make_chunk(doc, current_block, chunk_index))
                            chunk_index += 1
                            # overlap：保留上一个块的最后 overlap 字符作为开头
                            current_block = current_block[-self.overlap:] + "\n" + sent if len(current_block) > self.overlap else sent
                        else:
                            current_block = sent
            else:
                # 普通段落：尝试加入当前块
                if len(current_block) + len(para) + 2 <= self.chunk_size:
                    current_block = (current_block + "\n\n" + para).strip()
                else:
                    # 当前块已满，保存并开始新块（带 overlap）
                    if current_block:
                        chunks.append(self._make_chunk(doc, current_block, chunk_index))
                        chunk_index += 1
                        # 应用 overlap
                        current_block = current_block[-self.overlap:] + "\n\n" + para if len(current_block) > self.overlap else para
                    else:
                        current_block = para

        # 保存最后一块
        if current_block.strip():
            chunks.append(self._make_chunk(doc, current_block, chunk_index))

        return chunks

    def _split_sentences(self, text: str) -> List[str]:
        """按句子分割（支持中英文）"""
        # 中文句号/问号/感叹号，英文 .?!（排除缩写）
        sentences = re.split(r'(?<=[。！？.?!])\s*', text)
        return [s.strip() for s in sentences if s.strip()]

    def _make_chunk(self, doc: Document, text: str, idx: int) -> Chunk:
        """构建 Chunk 对象"""
        return Chunk(
            chunk_id=f"{doc.doc_id}_c{idx}",
            doc_id=doc.doc_id,
            text=text,
            chunk_index=idx,
            metadata={
                **doc.metadata,
                "doc_id": doc.doc_id,
                "chunk_index": idx,
                "char_count": len(text)
            }
        )


class WeChatChunker(TextChunker):
    """微信聊天记录专用分块器"""
    
    def __init__(self, chunk_size: int = 500, overlap: int = 50, time_window_minutes: int = 5):
        """
        初始化微信聊天记录分块器
        
        Args:
            chunk_size: 分块大小（字符数）
            overlap: 重叠字符数
            time_window_minutes: 时间窗口（分钟），同一发言人在此窗口内的消息合并
        """
        super().__init__(chunk_size, overlap)
        self.time_window_minutes = time_window_minutes
    
    def chunk_wechat_md(self, md_path: str) -> List[Chunk]:
        """
        分块微信聊天记录 Markdown 文件
        
        策略：
        - 同一话题的多轮对话合并为一个 chunk（直到 chunk_size 溢出或空闲超时）
        - 每个 chunk 带有发言人和时间范围元数据
        - 跳过无意义内容（纯 [图片] / [链接/文件] 等）
        """
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        chunks = []
        chunk_index = 0
        
        # 解析所有消息
        messages = self._parse_messages(content)
        if not messages:
            return chunks
        
        # 按主题窗口分组
        group = []
        for msg in messages:
            text = msg['text']
            if not text or text.strip() in ('[图片]', '[链接/文件]', ''):
                continue
            
            need_new_group = False
            if not group:
                need_new_group = True
            else:
                prev = group[-1]
                gap_minutes = self._time_diff(prev['timestamp'], msg['timestamp'])
                if gap_minutes > self.time_window_minutes:
                    need_new_group = True
                elif sum(len(m['raw']) for m in group) + len(msg['raw']) > self.chunk_size:
                    need_new_group = True
            
            if need_new_group and group:
                chunk = self._build_group_chunk(md_path, group, chunk_index)
                if chunk:
                    chunks.append(chunk)
                    chunk_index += 1
                group = []
            
            group.append(msg)
        
        if group:
            chunk = self._build_group_chunk(md_path, group, chunk_index)
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def _parse_messages(self, content: str) -> List[Dict]:
        """解析标准格式的微信聊天记录"""
        messages = []
        current_msg = None
        
        for line in content.split('\n'):
            match = re.match(r'\*\*(.+?)\*\* \((.+?)\):', line)
            if match:
                if current_msg:
                    messages.append(current_msg)
                current_msg = {
                    'speaker': match.group(1),
                    'timestamp': match.group(2),
                    'raw': line,
                    'text': ''
                }
            elif current_msg is not None:
                stripped = line.strip()
                if stripped:
                    current_msg['text'] += stripped
                    current_msg['raw'] += '\n' + line
        
        if current_msg:
            messages.append(current_msg)
        
        return messages
    
    def _build_group_chunk(self, md_path: str, messages: List[Dict], idx: int) -> Optional[Chunk]:
        """将一组消息构建为一个 Chunk"""
        if not messages:
            return None
        
        text = '\n'.join(m['raw'] for m in messages)
        speakers = list(dict.fromkeys(m['speaker'] for m in messages))
        time_start = messages[0]['timestamp']
        time_end = messages[-1]['timestamp']
        
        speakers_str = '、'.join(speakers[:4])
        if len(speakers) > 4:
            speakers_str += f' 等{len(speakers)}人'
        
        # 生成基于文件路径的唯一 doc_id（避免不同文档的 chunk 互相覆盖 ChromaDB）
        chunk_doc_id = hashlib.md5(md_path.encode()).hexdigest()[:12]
        return Chunk(
            chunk_id=f"wechat_{chunk_doc_id}_c{idx}",
            doc_id=chunk_doc_id,
            text=text,
            chunk_index=idx,
            metadata={
                "chunk_index": idx,
                "char_count": len(text),
                "speaker": speakers_str,
                "speakers": ','.join(speakers),
                "timestamp": time_start,
                "time_start": time_start,
                "time_end": time_end,
                "msg_count": len(messages),
                "source": "wechat"
            }
        )
    
    def _time_diff(self, t1_str: str, t2_str: str) -> int:
        """计算时间差（分钟）"""
        if not t1_str or not t2_str:
            return 0
        try:
            t1 = datetime.strptime(t1_str, '%Y-%m-%d %H:%M:%S')
            t2 = datetime.strptime(t2_str, '%Y-%m-%d %H:%M:%S')
            return abs((t2 - t1).total_seconds()) / 60
        except:
            return 0

    def extract_messages(self, md_path: str, chunk_results: List[Chunk]) -> List[Dict]:
        """从微信聊天记录中提取单条消息（方案 C：消息级索引）
        
        Args:
            md_path: 微信聊天记录 Markdown 文件路径
            chunk_results: 已分块的 Chunk 列表（用于建立消息→chunk 的映射）
        
        Returns:
            消息列表，每项格式:
            {
                'id': 'msg_{md5_hash}',
                'text': '消息文本',
                'metadata': {
                    'filename': '',
                    'sender': '发言人',
                    'time': '时间戳',
                    'chunk_id': '所属 chunk 的 ID',
                    'category': ''
                }
            }
        """
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 解析所有消息
        raw_messages = self._parse_messages(content)
        if not raw_messages:
            return []

        # 为每条消息找到所属的 chunk
        # 构建 chunk 索引：chunk 的文本覆盖范围
        chunk_map = []
        for chunk in chunk_results:
            first_line = chunk.text.split('\n')[0] if chunk.text else ''
            last_line = chunk.text.split('\n')[-1] if chunk.text else ''
            chunk_map.append({
                'chunk_id': chunk.chunk_id,
                'chunk': chunk,
                'first': first_line,
                'last': last_line
            })

        messages = []
        filename = Path(md_path).name
        chunk_doc_id = hashlib.md5(md_path.encode()).hexdigest()[:12]

        for msg_idx, msg in enumerate(raw_messages):
            msg_raw = msg['raw']
            msg_text = msg['text']
            if not msg_text or msg_text.strip() in ('[图片]', '[链接/文件]', ''):
                continue

            # 找到所属 chunk：通过消息 raw 文本在 chunk 中出现的位置
            assigned_chunk_id = None
            for cm in chunk_map:
                if msg_raw in cm['chunk'].text:
                    assigned_chunk_id = cm['chunk_id']
                    break

            if assigned_chunk_id is None:
                continue

            # 生成唯一消息 ID（用序号而非 hash，避免碰撞）
            msg_id = f"msg_{chunk_doc_id}_{msg_idx}"

            # 消息文本只包含本条消息（不包含上下文），保证向量检索精准
            short_text = f"{msg['speaker']} ({msg['timestamp']}): {msg_text}"

            messages.append({
                'id': msg_id,
                'text': short_text,
                'metadata': {
                    'filename': filename,
                    'sender': msg['speaker'],
                    'time': msg['timestamp'],
                    'chunk_id': assigned_chunk_id,
                    'category': ''  # 由调用方填充
                }
            })

        return messages


# ========== 3. Embedding 模型封装 ==========

class Embedder:
    """
    Embedding 模型封装

    支持：
    1. Ollama 本地模型（nomic-embed-text）- 首选
    2. HuggingFace 本地模型（text2vec-base-chinese 等）- 备选
    """

    def __init__(self, provider: str = "ollama", model: str = "nomic-embed-text"):
        self.provider = provider
        self.model = model
        self._session = None
        self._hf_model = None
        self._base_url = "http://127.0.0.1:11434"

    def _get_session(self):
        """获取 HTTP Session"""
        if self._session is None:
            import requests
            self._session = requests.Session()
        return self._session

    def embed_texts(self, texts: List[str]) -> List[List[float]]:
        """将多个文本向量化"""
        if self.provider == "ollama":
            return self._embed_ollama_batch(texts)
        elif self.provider == "huggingface":
            return self._embed_huggingface(texts)
        else:
            raise ValueError(f"不支持的 provider: {self.provider}")

    def _embed_ollama_batch(self, texts: List[str]) -> List[List[float]]:
        """通过 Ollama API 批量获取 embeddings"""
        session = self._get_session()
        embeddings = []

        for text in texts:
            try:
                response = session.post(
                    f"{self._base_url}/api/embeddings",
                    json={"model": self.model, "prompt": text},
                    timeout=60,
                    proxies={'http': None, 'https': None}
                )
                if response.status_code != 200:
                    logger.error(f"Ollama embedding 失败: {response.status_code} - {response.text}")
                    # 返回零向量作为降级
                    embeddings.append([0.0] * 768)
                    continue
                data = response.json()
                embedding = data.get('embedding', [])
                if not embedding:
                    embeddings.append([0.0] * 768)
                else:
                    embeddings.append(embedding)
            except Exception as e:
                logger.error(f"Ollama embedding 请求异常: {e}")
                embeddings.append([0.0] * 768)

        return embeddings

    def _embed_huggingface(self, texts: List[str]) -> List[List[float]]:
        """使用 HuggingFace 本地模型"""
        if self._hf_model is None:
            try:
                from sentence_transformers import SentenceTransformer
                logger.info("正在加载 HuggingFace Embedding 模型...")
                self._hf_model = SentenceTransformer('shenglOL/text2vec-base-chinese')
            except ImportError:
                raise ImportError("请先安装 sentence-transformers: pip install sentence-transformers")

        embeddings = self._hf_model.encode(texts, normalize_embeddings=True)
        return embeddings.tolist()

    def embed_query(self, query: str) -> List[float]:
        """向量化单个查询"""
        return self.embed_texts([query])[0]


# ========== 4. ChromaDB 向量存储 ==========

class VectorStore:
    """ChromaDB 向量存储封装
    
    双集合架构（方案 C）：
    - `knowledge_base`（主集合）：chunk 级存储，用于生成回答时的上下文
    - `knowledge_base_messages`（消息集合）：消息级存储，用于精准检索
    """

    def __init__(self, collection_name: str = "knowledge_base",
                 msg_collection_name: str = "knowledge_base_messages"):
        import chromadb
        from chromadb.config import Settings

        # 确保路径存在
        CHROMA_PATH.mkdir(parents=True, exist_ok=True)

        self.client = chromadb.PersistentClient(
            path=str(CHROMA_PATH),
            settings=Settings(anonymized_telemetry=False)
        )
        # 主集合：chunk 级
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            metadata={"description": "qwen3-asr 知识库 — Chunk 级（上下文用）"}
        )
        # 消息集合：消息级（精准检索用）
        self.msg_collection = self.client.get_or_create_collection(
            name=msg_collection_name,
            metadata={"description": "qwen3-asr 知识库 — 消息级（检索用）"}
        )
        logger.info(f"✅ ChromaDB 初始化完成")
        logger.info(f"   - Chunk 集合 '{collection_name}': {self.collection.count()} 块")
        logger.info(f"   - 消息集合 '{msg_collection_name}': {self.msg_collection.count()} 条")

    def add_chunks(self, chunks: List[Chunk], embeddings: List[List[float]]):
        """批量添加文本块"""
        if not chunks:
            return

        ids = [c.chunk_id for c in chunks]
        texts = [c.text for c in chunks]
        metadatas = [c.metadata for c in chunks]

        # 批量添加，每批 100 个
        batch_size = 100
        for i in range(0, len(chunks), batch_size):
            batch_ids = ids[i:i + batch_size]
            batch_texts = texts[i:i + batch_size]
            batch_metas = metadatas[i:i + batch_size]
            batch_embs = embeddings[i:i + batch_size]

            self.collection.add(
                ids=batch_ids,
                embeddings=batch_embs,
                documents=batch_texts,
                metadatas=batch_metas
            )

        logger.info(f"📦 已添加 {len(chunks)} 个文本块到向量库")

    def search(self, query_embedding: List[float], top_k: int = 5,
               where: Optional[Dict] = None) -> List[SearchHit]:
        """
        语义检索

        Args:
            query_embedding: 查询向量
            top_k: 返回前 N 条
            where: 元数据过滤条件

        Returns:
            检索结果列表
        """
        try:
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                where=where,
                include=["documents", "metadatas", "distances"]
            )

            hits = []
            if results and results['ids'] and len(results['ids']) > 0:
                for i in range(len(results['ids'][0])):
                    dist = results['distances'][0][i]
                    hits.append(SearchHit(
                        chunk_id=results['ids'][0][i],
                        text=results['documents'][0][i],
                        metadata=results['metadatas'][0][i],
                        distance=dist,
                        score=max(0.0, 1.0 - dist)  # 余弦距离转相似度
                    ))

            return hits

        except Exception as e:
            logger.error(f"向量检索失败: {e}")
            return []

    def delete_by_doc_id(self, doc_id: str):
        """删除某个文档的所有块"""
        try:
            results = self.collection.get(where={"doc_id": doc_id})
            if results and results['ids']:
                self.collection.delete(ids=results['ids'])
                logger.info(f"🗑️ 已删除文档 {doc_id} 的 {len(results['ids'])} 个块")
        except Exception as e:
            logger.error(f"删除文档块失败: {e}")

    def add_messages(self, messages: List[Dict], embeddings: List[List[float]]):
        """批量添加消息级索引（用于精准检索）
        
        Args:
            messages: 消息列表，每项包含 id, text, metadata
            embeddings: 对应的向量列表
        """
        if not messages:
            return

        ids = [m['id'] for m in messages]
        texts = [m['text'] for m in messages]
        metadatas = [m['metadata'] for m in messages]

        batch_size = 100
        for i in range(0, len(messages), batch_size):
            self.msg_collection.add(
                ids=ids[i:i + batch_size],
                embeddings=embeddings[i:i + batch_size],
                documents=texts[i:i + batch_size],
                metadatas=metadatas[i:i + batch_size]
            )

        logger.info(f"📋 已添加 {len(messages)} 条消息索引")

    def search_messages(self, query_embedding: List[float], top_k: int = 10,
                        where: Optional[Dict] = None) -> List[SearchHit]:
        """消息级检索
        
        搜索消息集合，返回每条消息的命中结果。
        """
        try:
            results = self.msg_collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                where=where,
                include=["documents", "metadatas", "distances"]
            )

            hits = []
            if results and results['ids'] and len(results['ids']) > 0:
                for i in range(len(results['ids'][0])):
                    dist = results['distances'][0][i]
                    hits.append(SearchHit(
                        chunk_id=results['ids'][0][i],
                        text=results['documents'][0][i],
                        metadata=results['metadatas'][0][i],
                        distance=dist,
                        score=max(0.0, 1.0 - dist)
                    ))

            return hits
        except Exception as e:
            logger.error(f"消息级检索失败: {e}")
            return []

    def get_chunk_by_id(self, chunk_id: str) -> Optional[SearchHit]:
        """通过 chunk_id 获取完整的 chunk 文本（用于上下文）"""
        try:
            results = self.collection.get(
                ids=[chunk_id],
                include=["documents", "metadatas"]
            )
            if results and results['ids'] and len(results['ids']) > 0:
                return SearchHit(
                    chunk_id=results['ids'][0],
                    text=results['documents'][0],
                    metadata=results['metadatas'][0],
                    distance=0.0,
                    score=1.0
                )
        except Exception as e:
            pass
        return None

    def delete_by_filename(self, filename: str, category: str = None) -> int:
        """删除指定文件名的所有块和消息（用于去重）"""
        total = 0
        for coll, label in [(self.collection, '块'), (self.msg_collection, '消息')]:
            try:
                results = coll.get(include=["metadatas"])
                if not results or not results['metadatas']:
                    continue

                ids_to_delete = []
                for i, meta in enumerate(results['metadatas']):
                    fn = meta.get('filename', '')
                    cat = meta.get('category', '')

                    fn_base = fn.replace('_raw.standard.md', '').replace('.standard.md', '').replace('.md', '')
                    name_match = (fn == filename or filename in fn or filename in fn_base or fn_base == filename)

                    if not name_match:
                        continue
                    if category is not None and cat != category:
                        continue

                    ids_to_delete.append(results['ids'][i])

                if ids_to_delete:
                    coll.delete(ids=ids_to_delete)
                    label_full = f"'{filename}'"
                    if category:
                        label_full += f" (category={category})"
                    logger.info(f"🗑️ 已删除 {label_full} 的 {len(ids_to_delete)} 个{label}")
                    total += len(ids_to_delete)
            except Exception as e:
                logger.error(f"删除 {label} 失败: {e}")
        return total

    def count(self) -> int:
        """向量库中的总块数"""
        return self.collection.count()

    def msg_count(self) -> int:
        """消息集合中的总消息数"""
        return self.msg_collection.count()

    def get_stats(self) -> Dict:
        """获取知识库统计信息（含消息级）"""
        stats = {
            "total_chunks": self.collection.count(),
            "total_messages": self.msg_collection.count(),
            "total_docs": 0,
            "categories": []
        }
        try:
            results = self.collection.get(include=["metadatas"])
            if results and results['metadatas']:
                seen_docs = set()
                seen_cats = set()
                for meta in results['metadatas']:
                    doc_id = meta.get('doc_id', '')
                    cat = meta.get('category', '默认')
                    if doc_id:
                        seen_docs.add(doc_id)
                    seen_cats.add(cat)
                stats['total_docs'] = len(seen_docs)
                stats['categories'] = list(seen_cats)
        except Exception as e:
            logger.error(f"获取统计信息失败: {e}")
        return stats

    def get_all_docs(self) -> List[Dict]:
        """获取所有文档的元数据"""
        try:
            results = self.collection.get(include=["metadatas"])
            if not results or not results['metadatas']:
                return []

            # 按 doc_id 分组
            docs_map = {}
            for meta in results['metadatas']:
                doc_id = meta.get('doc_id', '')
                if doc_id and doc_id not in docs_map:
                    docs_map[doc_id] = {
                        "doc_id": doc_id,
                        "filename": meta.get('filename', ''),
                        "source": meta.get('source', ''),
                        "category": meta.get('category', '默认'),
                        "loaded_at": meta.get('loaded_at', ''),
                        "chunk_count": 0
                    }
                if doc_id:
                    docs_map[doc_id]['chunk_count'] += 1

            return list(docs_map.values())
        except Exception as e:
            logger.error(f"获取文档列表失败: {e}")
            return []

    def clear_all(self):
        """清空所有数据（谨慎使用）"""
        try:
            self.client.delete_collection(name=self.collection.name)
            self.collection = self.client.get_or_create_collection(
                name=self.collection.name,
                metadata={"description": "qwen3-asr 知识库"}
            )
            logger.info("🗑️ 已清空知识库")
        except Exception as e:
            logger.error(f"清空知识库失败: {e}")


# ========== 5. RAG 检索链 ==========

class RAGChain:
    """
    RAG (Retrieval Augmented Generation) 检索链

    流程：
    1. 接收用户问题
    2. 向量化问题（Embedding）
    3. 从向量库检索相关块（top_k）
    4. 将检索结果注入提示词
    5. 调用 LLM 生成回答
    """

    def __init__(self, embedder: Embedder, vectorstore: VectorStore, summarizer: Any):
        self.embedder = embedder
        self.vectorstore = vectorstore
        self.summarizer = summarizer  # LongTextSummarizer 实例

    def _retrieve(self, question: str, top_k: int = 5,
                 where: Optional[Dict] = None) -> List[SearchHit]:
        """
        检索阶段（方案 C：消息级检索 + Chunk 上下文）
        
        流程：
        1. 先在消息集合中做精准检索（找到最相关的单条消息）
        2. 按 chunk_id 去重，取父 chunk 的完整上下文
        3. 如果消息集合为空（旧数据等），回退到 chunk 级检索
        """
        try:
            query_embedding = self.embedder.embed_query(question)
        except Exception as e:
            logger.error(f"问题向量化失败: {e}")
            return []

        # ---- 方案 C：消息级检索 ----
        msg_count = self.vectorstore.msg_collection.count()
        if msg_count > 0:
            return self._retrieve_via_messages(query_embedding, top_k, where)

        # ---- 回退：传统 chunk 级检索 ----
        logger.info("ℹ️ 消息集合为空，回退到 chunk 级检索")
        hits = self.vectorstore.search(
            query_embedding=query_embedding,
            top_k=top_k,
            where=where
        )
        return hits

    def _retrieve_via_messages(self, query_embedding: List[float],
                                top_k: int = 5,
                                where: Optional[Dict] = None) -> List[SearchHit]:
        """消息级检索 → 按 chunk 去重 → 取完整上下文"""
        # 1. 搜更多消息，确保覆盖足够的 chunk
        search_k = max(top_k * 3, 20)
        msg_hits = self.vectorstore.search_messages(
            query_embedding=query_embedding,
            top_k=search_k,
            where=where
        )

        if not msg_hits:
            return []

        # 2. 按 chunk_id 去重，保留匹配度最高的消息信息
        chunk_best = {}  # chunk_id → {best_msg, chunk_hit}
        for hit in msg_hits:
            chunk_id = hit.metadata.get('chunk_id', '')
            if not chunk_id:
                continue

            if chunk_id not in chunk_best or hit.score > chunk_best[chunk_id]['score']:
                # 获取完整 chunk 文本
                chunk_hit = self.vectorstore.get_chunk_by_id(chunk_id)
                if chunk_hit:
                    chunk_best[chunk_id] = {
                        'best_msg': hit,
                        'chunk_hit': chunk_hit,
                        'score': hit.score
                    }

        if not chunk_best:
            return []

        # 3. 按匹配度排序，取 top_k 个 chunk
        sorted_chunks = sorted(chunk_best.values(), key=lambda x: -x['score'])
        sorted_chunks = sorted_chunks[:top_k]

        # 4. 构建返回：用 chunk 文本 + 消息级元数据
        result = []
        for item in sorted_chunks:
            chunk_hit = item['chunk_hit']
            best_msg = item['best_msg']

            # 在 chunk 元数据中注入消息级信息
            enriched_meta = {
                **chunk_hit.metadata,
                '_matched_sender': best_msg.metadata.get('sender', ''),
                '_matched_time': best_msg.metadata.get('time', ''),
                '_matched_text': best_msg.text[:120],  # 命中的消息预览
            }

            result.append(SearchHit(
                chunk_id=chunk_hit.chunk_id,
                text=chunk_hit.text,
                metadata=enriched_meta,
                distance=best_msg.distance,
                score=best_msg.score
            ))

        logger.info(f"🔍 消息级检索: {len(msg_hits)} hits → "
                    f"{len(chunk_best)} unique chunks → {len(result)} top")
        return result

    def _generate(self, question: str, hits: List[SearchHit],
                  model_id: Optional[str] = None) -> str:
        """仅执行生成阶段，基于给定的 hits 生成回答"""
        if not hits:
            return "抱歉，知识库中没有找到与您问题相关的内容。"

        context = self._build_context(hits)
        prompt = self._build_prompt(question, context)

        system_prompt = """你是一个知识库问答助手。你的任务是根据提供的参考资料回答用户的问题。

规则：
1. 必须严格按照提供的参考资料回答，不要编造信息
2. 如果参考资料中有相关内容，引用时注明来源（文件名）
3. 如果参考资料中没有相关信息，明确告知用户
4. 回答要准确、完整、简洁
5. 优先使用参考资料中的原文表述，必要时进行总结"""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ]

        try:
            answer = self.summarizer.chat(messages, max_new_tokens=2048)
        except Exception as e:
            logger.error(f"LLM 生成失败: {e}")
            answer = f"抱歉，生成回答时出现错误：{str(e)}"

        return answer

    def query(self, question: str, top_k: int = 5,
              filter_category: Optional[str] = None,
              where: Optional[Dict] = None) -> Tuple[str, List[SearchHit]]:
        """
        完整 RAG 查询（检索 + 生成）

        Args:
            where: 直接传入元数据过滤条件（比 filter_category 更灵活）
        """
        # 1. 检索
        if where is None and filter_category:
            where = {"category": filter_category}

        hits = self._retrieve(question, top_k, where)

        if not hits:
            return "抱歉，知识库中没有找到与您问题相关的内容。", []

        # 2. 生成
        answer = self._generate(question, hits)

        return answer, hits

    def _build_context(self, hits: List[SearchHit]) -> str:
        """将检索结果构建为上下文"""
        context_parts = []
        for i, hit in enumerate(hits, 1):
            filename = hit.metadata.get('filename', '未知来源')
            source = hit.metadata.get('source', '')
            context_parts.append(
                f"【参考资料 {i}】（来源：{filename}）\n{hit.text}"
            )
        return "\n\n---\n\n".join(context_parts)

    def _build_prompt(self, question: str, context: str) -> str:
        """构建用户提示词"""
        return f"""【用户问题】
{question}

【参考资料】
以下是从知识库中检索到的相关资料，请仔细阅读后回答：

{context}

请根据以上参考资料回答用户问题。如果资料中有明确答案，给出准确回答并注明参考来源。"""


# ========== 6. 全局单例与管理 ==========

# 全局实例
_loader: Optional[DocumentLoader] = None
_chunker: Optional[TextChunker] = None
_embedder: Optional[Embedder] = None
_vectorstore: Optional[VectorStore] = None
_rag_chain: Optional[RAGChain] = None
_summarizer_ref: Any = None


def init_knowledge_base(summarizer: Any = None,
                        chunk_size: int = 500,
                        overlap: int = 50,
                        embed_provider: str = "ollama",
                        embed_model: str = "nomic-embed-text") -> bool:
    """
    初始化知识库模块

    Args:
        summarizer: LongTextSummarizer 实例
        chunk_size: 分块大小（字符数）
        overlap: 重叠字符数
        embed_provider: Embedding 提供者 ("ollama" 或 "huggingface")
        embed_model: Embedding 模型名称

    Returns:
        是否初始化成功
    """
    global _loader, _chunker, _embedder, _vectorstore, _rag_chain, _summarizer_ref, _last_init_error

    try:
        logger.info("🔧 正在初始化知识库模块...")
        logger.info(f"   - 分块大小: {chunk_size} 字符, 重叠: {overlap} 字符")
        logger.info(f"   - Embedding: {embed_provider} / {embed_model}")

        _summarizer_ref = summarizer
        _loader = DocumentLoader()
        _chunker = TextChunker(chunk_size=chunk_size, overlap=overlap)
        _embedder = Embedder(provider=embed_provider, model=embed_model)
        _vectorstore = VectorStore(collection_name="knowledge_base")
        if summarizer is not None:
            _rag_chain = RAGChain(_embedder, _vectorstore, summarizer)
        else:
            _rag_chain = None

        logger.info("✅ 知识库模块初始化完成")
        return True

    except Exception as e:
        logger.error(f"❌ 知识库初始化失败: {e}")
        # 记录详细堆栈到 stderr，方便排查
        import traceback
        traceback.print_exc()
        # 把错误存到全局变量，让 API 层可以读取
        global _last_init_error
        _last_init_error = str(e)
        return False


def get_last_init_error() -> str:
    """获取最后一次初始化失败的具体错误"""
    global _last_init_error
    return _last_init_error or "未知错误"


def get_loader() -> DocumentLoader:
    if _loader is None:
        raise RuntimeError("知识库未初始化，请先调用 init_knowledge_base()")
    return _loader


def get_chunker() -> TextChunker:
    if _chunker is None:
        raise RuntimeError("知识库未初始化，请先调用 init_knowledge_base()")
    return _chunker


def get_embedder() -> Embedder:
    if _embedder is None:
        raise RuntimeError("知识库未初始化，请先调用 init_knowledge_base()")
    return _embedder


def get_vectorstore() -> VectorStore:
    if _vectorstore is None:
        raise RuntimeError("知识库未初始化，请先调用 init_knowledge_base()")
    return _vectorstore


def get_rag_chain() -> RAGChain:
    if _rag_chain is None:
        raise RuntimeError("知识库未初始化，请先调用 init_knowledge_base()")
    return _rag_chain


def delete_by_filename(filename: str, category: str = None) -> int:
    """删除指定文件名的所有块（用于去重）
    
    Args:
        filename: 文件名（如群名、文档名）
        category: 可选，限定只删除特定 category 的数据
    
    Returns:
        删除的块数
    """
    if _vectorstore is None:
        raise RuntimeError("知识库未初始化，请先调用 init_knowledge_base()")
    return _vectorstore.delete_by_filename(filename, category=category)


def is_initialized() -> bool:
    """检查知识库是否已初始化"""
    return _rag_chain is not None


# ========== 辅助函数 ==========

def _is_wechat_record(file_path: str) -> bool:
    """检测文件是否为微信聊天记录（标准格式）"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            for i, line in enumerate(f):
                if i > 20:
                    break
                if re.match(r'\*\*.+?\*\* \(.+?\):', line):
                    return True
        return False
    except:
        return False


def index_document(file_path: str, category: str = "默认") -> Dict:
    """
    一键索引文档（加载 → 分块 → 向量化 → 存储）

    Args:
        file_path: 文档路径
        category: 分类标签

    Returns:
        索引结果信息
    """
    loader = get_loader()
    embedder = get_embedder()
    vectorstore = get_vectorstore()

    # 1. 加载文档
    doc = loader.load(file_path)
    doc.metadata['category'] = category

    # 保存原始文件到知识库目录
    dest_dir = DOCS_PATH / category.replace("/", "_")
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest_path = dest_dir / doc.filename

    # 复制文件
    import shutil
    shutil.copy2(file_path, dest_path)
    doc.metadata['stored_path'] = str(dest_path)

    # 2. 分块（微信聊天记录用专用分块器）
    if _is_wechat_record(file_path):
        logger.info("💬 检测到微信聊天记录，使用 WeChatChunker")
        wc = WeChatChunker(chunk_size=800, overlap=50, time_window_minutes=10)
        chunks = wc.chunk_wechat_md(file_path)
        # 注入 doc_id 和 category 到每个 chunk 的元数据
        for i, c in enumerate(chunks):
            c.doc_id = doc.doc_id
            c.metadata['doc_id'] = doc.doc_id
            c.metadata['category'] = category
            c.metadata['filename'] = doc.filename
            # 修复: chunk_id 必须包含 doc_id 以避免不同文档的 chunk 互相覆盖
            c.chunk_id = f"wechat_{doc.doc_id}_c{i}"
        doc.chunk_count = len(chunks)
    else:
        chunker = get_chunker()
        chunks = chunker.chunk(doc)
        doc.chunk_count = len(chunks)

    logger.info(f"📄 文档 '{doc.filename}' 已分块: {len(chunks)} 个块")

    # 3. 向量化（chunk 级）
    texts = [c.text for c in chunks]
    embeddings = embedder.embed_texts(texts)

    # 4. 存入向量库（chunk 集合）
    vectorstore.add_chunks(chunks, embeddings)

    # 方案 C：如果是微信聊天记录，同时创建消息级索引
    message_entries = []
    if _is_wechat_record(file_path):
        wc = WeChatChunker(chunk_size=800, overlap=50, time_window_minutes=10)
        message_entries = wc.extract_messages(file_path, chunks)
        for m in message_entries:
            m['metadata']['category'] = category
        
    if message_entries:
        msg_texts = [m['text'] for m in message_entries]
        msg_embeddings = embedder.embed_texts(msg_texts)
        vectorstore.add_messages(message_entries, msg_embeddings)
        logger.info(f"📋 同时创建了 {len(message_entries)} 条消息级索引")

    # 5. 更新索引文件
    _update_index_file(doc)

    return {
        "doc_id": doc.doc_id,
        "filename": doc.filename,
        "chunk_count": len(chunks),
        "category": category
    }


def _update_index_file(doc: Document):
    """更新文档索引文件"""
    index_file = KB_ROOT / "index.json"
    index = {}

    if index_file.exists():
        try:
            index = json.loads(index_file.read_text(encoding='utf-8'))
        except Exception:
            index = {}

    index[doc.doc_id] = {
        "filename": doc.filename,
        "metadata": doc.metadata,
        "chunk_count": doc.chunk_count,
        "indexed_at": doc.metadata.get('loaded_at', '')
    }

    index_file.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding='utf-8')
    logger.info(f"📋 索引文件已更新: {index_file}")
